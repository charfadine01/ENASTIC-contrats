"""Générateur de contrats Word/PDF — porté depuis l'app Flask ENASTIC."""

import hashlib
import json
import logging
import os
import shutil
import subprocess
import tempfile
import uuid
from copy import deepcopy
from datetime import datetime
from io import BytesIO
from typing import Optional

import qrcode
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches

logger = logging.getLogger(__name__)


def _detect_libreoffice() -> Optional[str]:
    candidates = [
        "/opt/homebrew/bin/soffice",
        "/usr/local/bin/soffice",
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
        "soffice",
    ]
    for path in candidates:
        if shutil.which(path) or os.path.exists(path):
            return path
    return None


LIBREOFFICE_PATH = _detect_libreoffice()
PDF_AVAILABLE = LIBREOFFICE_PATH is not None


class DocumentGeneratorError(Exception):
    pass


class ContractGenerator:
    GRADE_MAPPING = {
        "Professeur": "Professeur",
        "Maître de Conférences": "Maître de Conférences",
        "Maître Assistant": "Maître Assistant",
        "Assistant d'Université": "Assistant d'Université",
        "Assistant": "Assistant",
    }

    def __init__(self, template_path: str, output_folder: str):
        if not os.path.exists(template_path):
            raise DocumentGeneratorError(f"Template introuvable: {template_path}")
        self.template_path = template_path
        self.output_folder = output_folder
        os.makedirs(output_folder, exist_ok=True)

    @staticmethod
    def _set_cell_background(cell, color: str) -> None:
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), color)
        cell._element.get_or_add_tcPr().append(shading)

    @staticmethod
    def _replace_in_paragraph(paragraph, old: str, new: str) -> None:
        if old in paragraph.text:
            for run in paragraph.runs:
                if old in run.text:
                    run.text = run.text.replace(old, new)

    def _replace_in_table(self, table, old: str, new: str) -> None:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    self._replace_in_paragraph(para, old, new)

    def _highlight_grade_row(self, table, grade: str) -> None:
        target = self.GRADE_MAPPING.get(grade)
        if not target:
            return
        for i, row in enumerate(table.rows):
            if i == 0:
                continue
            if target in row.cells[0].text.strip():
                for cell in row.cells:
                    self._set_cell_background(cell, "FFFF00")
                break

    @staticmethod
    def _is_ecue_table(table) -> bool:
        if not table.rows:
            return False
        first = table.rows[0].cells[0].text
        return "Intitulé" in first or "ECUE" in first

    @staticmethod
    def _ensure_cell_borders(cell, color: str = "000000", size: str = "4") -> None:
        """Force des bordures noires fines sur les 4 côtés d'une cellule."""
        tc_pr = cell._element.get_or_add_tcPr()
        # Supprimer un éventuel tcBorders existant
        existing = tc_pr.find(qn("w:tcBorders"))
        if existing is not None:
            tc_pr.remove(existing)
        borders = OxmlElement("w:tcBorders")
        for side in ("top", "left", "bottom", "right"):
            elem = OxmlElement(f"w:{side}")
            elem.set(qn("w:val"), "single")
            elem.set(qn("w:sz"), size)
            elem.set(qn("w:space"), "0")
            elem.set(qn("w:color"), color)
            borders.append(elem)
        tc_pr.append(borders)

    def _fill_ecue_table(self, table, ecues: list[dict]) -> None:
        if not ecues or len(table.rows) <= 1:
            return

        # Sauvegarder le style de la ligne modèle (ligne 1)
        model_row = table.rows[1]
        cell_styles = [
            deepcopy(c._element.tcPr) if c._element.tcPr is not None else None
            for c in model_row.cells
        ]
        # Sauvegarder aussi le style de paragraphe pour préserver alignement/police
        para_styles = []
        for c in model_row.cells:
            if c.paragraphs and c.paragraphs[0]._element.pPr is not None:
                para_styles.append(deepcopy(c.paragraphs[0]._element.pPr))
            else:
                para_styles.append(None)

        # Supprimer toutes les lignes sauf l'en-tête
        for i in range(len(table.rows) - 1, 0, -1):
            row_elem = table.rows[i]._element
            row_elem.getparent().remove(row_elem)

        # Recréer une ligne par ECUE en réappliquant explicitement les bordures
        for ecue in ecues:
            new_row = table.add_row()
            values = [
                ecue.get("intitule", ""),
                str(ecue.get("heures_cm", "")),
                str(ecue.get("heures_td", "")),
                str(ecue.get("heures_tp", "")),
            ]
            for j, cell in enumerate(new_row.cells):
                cell.text = values[j] if j < len(values) else ""

                # Réappliquer tout le tcPr (largeur, alignement vertical, bordures)
                if j < len(cell_styles) and cell_styles[j] is not None:
                    existing_tcpr = cell._element.find(qn("w:tcPr"))
                    if existing_tcpr is not None:
                        cell._element.remove(existing_tcpr)
                    cell._element.insert(0, deepcopy(cell_styles[j]))

                # Garantir des bordures visibles (filet noir 0.5pt)
                self._ensure_cell_borders(cell)

                # Réappliquer l'alignement du paragraphe modèle
                if j < len(para_styles) and para_styles[j] is not None and cell.paragraphs:
                    para = cell.paragraphs[0]
                    existing_ppr = para._element.find(qn("w:pPr"))
                    if existing_ppr is not None:
                        para._element.remove(existing_ppr)
                    para._element.insert(0, deepcopy(para_styles[j]))

    @staticmethod
    def _hash_contract(data: dict, contract_uuid: str) -> str:
        payload = json.dumps(
            {
                "uuid": contract_uuid,
                "enseignant": data.get("nom_enseignant", ""),
                "annee_academique": data.get("annee_academique", ""),
                "grade": data.get("grade", ""),
                "timestamp": datetime.utcnow().isoformat(),
            },
            sort_keys=True,
        )
        return hashlib.sha256(payload.encode()).hexdigest()[:16]

    @staticmethod
    def _build_qr_image(data: dict, contract_uuid: str, verification_hash: str) -> BytesIO:
        qr_payload = {
            "type": "CONTRAT_ENASTIC",
            "uuid": contract_uuid,
            "enseignant": data.get("nom_enseignant", ""),
            "annee": data.get("annee_academique", ""),
            "grade": data.get("grade", ""),
            "date_generation": datetime.now().strftime("%d/%m/%Y %H:%M"),
            "hash": verification_hash,
        }
        qr = qrcode.QRCode(
            version=1,
            error_correction=qrcode.constants.ERROR_CORRECT_H,
            box_size=10,
            border=2,
        )
        qr.add_data(json.dumps(qr_payload, ensure_ascii=False))
        qr.make(fit=True)
        img = qr.make_image(fill_color="black", back_color="white")
        buffer = BytesIO()
        img.save(buffer, format="PNG")
        buffer.seek(0)
        return buffer

    @staticmethod
    def _build_signature_line(text_left: str, text_right: str, bold: bool = False) -> tuple:
        """
        Reconstruit un paragraphe avec :
          - texte gauche aligné au bord gauche
          - tabstop droit à 14 cm pour aligner texte droit (laisse une marge à droite)
          - bordures verticales: aucune ; espacement minimal
        Renvoie (pPr_xml, runs_xml) en string pour insertion via OxmlElement.
        """
        # Pas utilisé directement — on construit via API python-docx ci-dessous.
        return text_left, text_right

    SIGNATURE_FONT = "Times New Roman"
    SIGNATURE_SIZE_PT = 12

    def _style_signature_run(self, run, bold: bool, underline: bool) -> None:
        from docx.shared import Pt

        run.font.name = self.SIGNATURE_FONT
        run.font.size = Pt(self.SIGNATURE_SIZE_PT)
        # Forcer aussi les variantes east-asia / cs (sinon LibreOffice/Word peuvent garder Open Sans)
        rpr = run._element.get_or_add_rPr()
        # Supprimer un rFonts existant pour le réécrire proprement
        existing = rpr.find(qn("w:rFonts"))
        if existing is not None:
            rpr.remove(existing)
        rfonts = OxmlElement("w:rFonts")
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rfonts.set(qn(attr), self.SIGNATURE_FONT)
        rpr.insert(0, rfonts)
        if bold:
            run.bold = True
        if underline:
            run.underline = True

    def _reformat_signature_paragraph(
        self,
        para,
        left_text: str,
        right_text: str,
        bold: bool = False,
        underline: bool = False,
    ) -> None:
        """
        Vide un paragraphe et le reconstruit avec :
        - run gauche stylé
        - tab (sans style, pour que le souligné ne s'étende pas)
        - run droite stylé
        Police forcée Times New Roman 11 pt, tabstop droit à 14 cm.
        """
        from docx.enum.text import WD_TAB_ALIGNMENT
        from docx.shared import Cm

        for run in list(para.runs):
            run._element.getparent().remove(run._element)

        para.alignment = 0  # LEFT
        para.paragraph_format.left_indent = Cm(0)
        para.paragraph_format.first_line_indent = Cm(0)
        para.paragraph_format.right_indent = Cm(0)

        tab_stops = para.paragraph_format.tab_stops
        for ts in list(tab_stops):
            ts._element.getparent().remove(ts._element)
        tab_stops.add_tab_stop(Cm(14), WD_TAB_ALIGNMENT.RIGHT)

        run_left = para.add_run(left_text)
        self._style_signature_run(run_left, bold, underline)

        run_tab = para.add_run()
        run_tab.add_tab()
        # Pas de souligné/gras sur le tab → l'espace entre les 2 noms reste neutre

        run_right = para.add_run(right_text)
        self._style_signature_run(run_right, bold, underline)

    def _fix_signature_layout(self, doc: Document, vacataire_name: str, dg_name: str) -> None:
        """
        Reformate la zone de signature :
          - Ligne titres "Le Vacataire" / "Le Directeur Général de l'ENASTIC" → gras
          - Ligne noms vacataire / DG → gras + souligné
        Toujours sur une seule ligne grâce à un tabstop droit à 14 cm.
        """
        for para in doc.paragraphs:
            text = para.text
            if "Le Vacataire" in text and "Le Directeur Général" in text:
                self._reformat_signature_paragraph(
                    para,
                    "Le Vacataire",
                    "Le Directeur Général de l’ENASTIC",
                    bold=True,
                    underline=False,
                )
            elif (
                vacataire_name
                and dg_name
                and vacataire_name in text
                and dg_name in text
            ):
                self._reformat_signature_paragraph(
                    para,
                    vacataire_name,
                    dg_name,
                    bold=True,
                    underline=True,
                )

    def _embed_qr_code(self, doc: Document, qr_image: BytesIO) -> None:
        """
        Place un petit QR code (~8 mm) discret en bas à GAUCHE de la même page
        que les signataires. Pas d'espace ajouté : le QR suit naturellement.
        """
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Mm, Pt

        with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp:
            tmp.write(qr_image.getvalue())
            temp_path = tmp.name
        try:
            self._strip_trailing_blank_paragraphs(doc)

            qr_para = doc.add_paragraph()
            qr_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            qr_para.paragraph_format.left_indent = Mm(0)
            qr_para.paragraph_format.first_line_indent = Mm(0)
            qr_para.paragraph_format.space_before = Pt(12)  # ≈ 4 mm seulement
            qr_para.paragraph_format.space_after = Pt(0)
            run = qr_para.add_run()
            run.add_picture(temp_path, width=Mm(8))
        finally:
            try:
                os.unlink(temp_path)
            except OSError:
                pass

    @staticmethod
    def _set_bottom_margin(doc: Document, cm_value: float, footer_cm: float | None = None) -> None:
        """
        Définit la marge basse de chaque section à `cm_value` cm et,
        optionnellement, la distance du pied de page à `footer_cm`.
        """
        from docx.shared import Cm

        for section in doc.sections:
            section.bottom_margin = Cm(cm_value)
            if footer_cm is not None:
                section.footer_distance = Cm(footer_cm)

    @staticmethod
    def _strip_trailing_blank_paragraphs(doc: Document) -> None:
        """
        Supprime tous les paragraphes vides en fin de document, ainsi que
        les sauts de page (<w:br w:type="page"/>) résiduels, pour éviter
        que le QR code n'apparaisse seul sur une page suivante.
        """
        paragraphs = doc.paragraphs
        # Retirer les sauts de page explicites partout
        for para in paragraphs:
            for br in para._element.findall(qn("w:r") + "/" + qn("w:br")):
                if br.get(qn("w:type")) == "page":
                    br.getparent().remove(br)
        # Supprimer les paragraphes vides de la fin
        for para in reversed(paragraphs):
            if para.text.strip() == "" and not para._element.findall(qn("w:r") + "/" + qn("w:drawing")):
                para._element.getparent().remove(para._element)
            else:
                break

    @staticmethod
    def _tighten_spacing_around_articles(doc: Document) -> None:
        """
        Réduit l'espace vertical en supprimant les paragraphes vides qui précèdent
        Article 2, Article 3 et Article 5.
        Pour Article 3 (qui suit le tableau ECUE), on retire aussi un éventuel
        paragraphe vide placé juste après le tableau.
        """
        targets = ("Article 2", "Article 3", "Article 4", "Article 5")
        # On retravaille à chaque fois sur la liste fraîche car les indices changent
        for target in targets:
            paragraphs = doc.paragraphs
            idx = next(
                (i for i, p in enumerate(paragraphs) if p.text.strip().startswith(target)),
                None,
            )
            if idx is None:
                continue
            j = idx - 1
            while j >= 0 and paragraphs[j].text.strip() == "":
                empty_elem = paragraphs[j]._element
                empty_elem.getparent().remove(empty_elem)
                j -= 1
                paragraphs = doc.paragraphs  # rafraîchir après suppression
                idx = next(
                    (i for i, p in enumerate(paragraphs) if p.text.strip().startswith(target)),
                    None,
                )
                if idx is None:
                    break
                j = idx - 1
                # Bornes de sécurité
                if j < 0 or paragraphs[j].text.strip() != "":
                    break

    def _convert_to_pdf(self, docx_path: str) -> tuple[Optional[str], Optional[str]]:
        if not PDF_AVAILABLE:
            return None, "LibreOffice non disponible"
        try:
            output_dir = os.path.dirname(docx_path)
            cmd = [
                LIBREOFFICE_PATH,
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                output_dir,
                docx_path,
            ]
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
            if result.returncode != 0:
                return None, (result.stderr or result.stdout or "Erreur LibreOffice").strip()
            pdf_path = docx_path.replace(".docx", ".pdf")
            if not os.path.exists(pdf_path):
                return None, "PDF non créé"
            return pdf_path, None
        except subprocess.TimeoutExpired:
            return None, "Timeout (>30s)"
        except Exception as exc:
            return None, str(exc)

    def generate(self, data: dict) -> dict:
        doc = Document(self.template_path)
        grade = data.get("grade", "")
        annee = str(data.get("annee", ""))

        replacements = {
            "/2019": f"/{annee}" if annee else "/2019",
            "AAAA": data.get("nom_enseignant", ""),
            "MOUKHTAR HASSAN MAHAMAT": data.get("nom_enseignant", ""),
            "BBBB": grade,
            "CCCC": data.get("annee_academique", ""),
            "Dr HAGGAR BACHAR SALIM": data.get("directeur_general") or "Dr HAGGAR BACHAR SALIM",
            "052/PM/2000": data.get("arrete") or "052/PM/2000",
        }

        for paragraph in doc.paragraphs:
            for old, new in replacements.items():
                self._replace_in_paragraph(paragraph, old, new)

        # Reformater la zone de signature (Vacataire | Directeur Général)
        # pour garantir alignement sur une seule ligne, DG en retrait du bord droit
        vacataire_name = data.get("nom_enseignant", "")
        dg_name = data.get("directeur_general") or "Dr HAGGAR BACHAR SALIM"
        self._fix_signature_layout(doc, vacataire_name, dg_name)

        # Resserrer l'espace entre la fin de l'Article 4 et l'Article 5
        self._tighten_spacing_around_articles(doc)

        # Marges resserrées en bas pour garder le QR sur la même page
        # (marge basse 1,2 cm, pied de page 0,6 cm)
        self._set_bottom_margin(doc, 1.2, footer_cm=0.6)

        ecues = data.get("ecues", [])
        for table in doc.tables:
            if self._is_ecue_table(table):
                self._fill_ecue_table(table, ecues)
            else:
                for old, new in replacements.items():
                    self._replace_in_table(table, old, new)
                self._highlight_grade_row(table, grade)

        contract_uuid = str(uuid.uuid4())
        teacher_name = data.get("nom_enseignant", "Enseignant")
        academic_year = data.get("annee_academique", "")
        safe_name = teacher_name.replace(" ", "_").replace("/", "-")
        safe_year = academic_year.replace("/", "-")
        secure_filename = f"{contract_uuid}_Contrat_{safe_name}_{safe_year}.docx"
        display_filename = f"Contrat_{safe_name}_{safe_year}.docx"

        verification_hash = self._hash_contract(data, contract_uuid)
        qr_image = self._build_qr_image(data, contract_uuid, verification_hash)
        self._embed_qr_code(doc, qr_image)

        docx_path = os.path.join(self.output_folder, secure_filename)
        doc.save(docx_path)
        logger.info("Contrat généré uuid=%s hash=%s", contract_uuid, verification_hash)

        pdf_path = None
        pdf_filename = None
        if PDF_AVAILABLE:
            pdf_path, err = self._convert_to_pdf(docx_path)
            if err:
                logger.warning("PDF non généré: %s", err)
            elif pdf_path:
                pdf_filename = os.path.basename(pdf_path)

        return {
            "uuid": contract_uuid,
            "verification_hash": verification_hash,
            "docx_path": docx_path,
            "docx_filename": secure_filename,
            "display_filename": display_filename,
            "pdf_path": pdf_path,
            "pdf_filename": pdf_filename,
        }
